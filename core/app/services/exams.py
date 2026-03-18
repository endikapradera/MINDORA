from __future__ import annotations

import json
import re
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Literal

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.services.text_extract import extract_text_from_file
from app.services.query import retrieve_chunks
from app.services.llm import generate_text
from app.services.exam_validator import validate_exam_questions, score_solved_answer
from app.storage.branches import get_branch_path, branch_exists


def _exams_dir(branch: str) -> Path:
    return get_branch_path(branch) / "Exams"


def _simulations_dir(branch: str) -> Path:
    return _exams_dir(branch) / "simulations"


def _compact_context_texts(context_texts: list[str], max_chars: int = 4200) -> str:
    parts: list[str] = []
    total = 0
    for txt in context_texts:
        cleaned = re.sub(r"\s+", " ", txt).strip()
        if not cleaned:
            continue
        remaining = max_chars - total
        if remaining <= 0:
            break
        piece = cleaned[:remaining]
        parts.append(piece)
        total += len(piece) + 2
    return "\n\n".join(parts)


def _build_type_instructions(exam_type: Literal["test_simple", "test_multiple", "desarrollo", "mixto"]) -> str:
    if exam_type == "test_simple":
        return "Todas las preguntas deben ser tipo test de respuesta única con 4 opciones (A-D)."
    if exam_type == "test_multiple":
        return "Todas las preguntas deben ser tipo test de respuesta múltiple con 4 opciones (A-D) y al menos 2 correctas cuando aplique."
    if exam_type == "desarrollo":
        return "Todas las preguntas deben ser de desarrollo (sin opciones), con respuesta modelo clara y breve."
    return (
        "El examen debe ser mixto: combina preguntas tipo test simple, test múltiple y desarrollo. "
        "Asegura una distribución equilibrada."
    )


def _parse_generated_questions(raw: str) -> list[dict]:
    pattern = re.compile(r"(?ms)^###\s*Pregunta\s*(\d+)\s*$")
    matches = list(pattern.finditer(raw))
    if not matches:
        return []

    questions: list[dict] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(raw)
        block = raw[start:end].strip()
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]

        q_type = "desarrollo"
        statement = ""
        options: list[str] = []
        answer = ""
        explanation = ""

        in_options = False
        for line in lines:
            lowered = line.lower()
            if lowered.startswith("tipo:"):
                q_type = line.split(":", 1)[1].strip().lower()
                in_options = False
            elif lowered.startswith("enunciado:"):
                statement = line.split(":", 1)[1].strip()
                in_options = False
            elif lowered.startswith("opciones:"):
                in_options = True
            elif lowered.startswith("respuesta:"):
                answer = line.split(":", 1)[1].strip()
                in_options = False
            elif lowered.startswith("explicacion:") or lowered.startswith("explicación:"):
                explanation = line.split(":", 1)[1].strip()
                in_options = False
            elif in_options and re.match(r"^[A-D]\)", line):
                options.append(line)
            elif not statement:
                statement = line

        if statement and answer:
            questions.append(
                {
                    "number": len(questions) + 1,
                    "type": q_type,
                    "statement": statement,
                    "options": options,
                    "answer": answer,
                    "explanation": explanation,
                }
            )
    return questions


def _parse_fallback_questions(raw: str) -> list[dict]:
    splitter = re.compile(r"(?m)^\s*(\d{1,2})[\)\.:\-]\s+")
    matches = list(splitter.finditer(raw))
    if not matches:
        return []

    out: list[dict] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
        block = raw[start:end].strip()
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue

        statement = lines[0]
        options = [ln for ln in lines if re.match(r"^[A-D][\).]", ln)]

        answer_match = re.search(
            r"(?im)(respuesta|soluci[oó]n|correcta)\s*:\s*(.+)$",
            block,
        )
        answer = answer_match.group(2).strip() if answer_match else ""
        if not answer:
            continue

        q_type = "desarrollo"
        if options:
            if re.search(r"\b[ABCD]\s*[,y]\s*[ABCD]\b|\b[ABCD]\s*,\s*[ABCD]\b", answer, flags=re.IGNORECASE):
                q_type = "test_multiple"
            else:
                q_type = "test_simple"

        out.append(
            {
                "number": len(out) + 1,
                "type": q_type,
                "statement": statement,
                "options": options,
                "answer": answer,
                "explanation": "",
            }
        )
    return out


def _normalize_question_statement(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _dedupe_questions(questions: list[dict]) -> list[dict]:
    seen: set[str] = set()
    out: list[dict] = []
    for q in questions:
        key = _normalize_question_statement(q.get("statement", ""))
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(q)
    for i, q in enumerate(out, start=1):
        q["number"] = i
    return out


def _complete_missing_questions(
    *,
    topic: str,
    difficulty: str,
    exam_type: str,
    context_block: str,
    current_questions: list[dict],
    missing: int,
) -> list[dict]:
    if missing <= 0:
        return []

    existing_statements = "\n".join(f"- {q.get('statement', '')}" for q in current_questions[:40])
    completion_prompt = (
        "Completa el examen sin repetir preguntas ya existentes.\n"
        "Devuelve SOLO bloques en formato exacto:\n"
        "### Pregunta N\n"
        "Tipo: test_simple | test_multiple | desarrollo\n"
        "Enunciado: ...\n"
        "Opciones:\n"
        "A) ...\nB) ...\nC) ...\nD) ...\n"
        "Respuesta: ...\n"
        "Explicacion: ...\n"
        "Sin texto adicional fuera del formato.\n"
        f"Genera exactamente {missing} preguntas nuevas.\n"
        f"Tema: {topic}. Dificultad: {difficulty}. Tipo objetivo: {exam_type}.\n"
        f"Preguntas ya existentes (NO repetir):\n{existing_statements}\n\n"
        f"Contexto:\n{context_block}\n"
    )
    extra_raw = generate_text(completion_prompt, max_tokens=900, temperature=0.15)
    extra = _parse_generated_questions(extra_raw)
    if not extra:
        extra = _parse_fallback_questions(extra_raw)
    return extra


def _render_exam_statement(questions: list[dict], topic: str, difficulty: str, exam_type: str) -> str:
    lines = [
        f"EXAMEN - {topic}",
        f"Dificultad: {difficulty}",
        f"Tipo: {exam_type}",
        "",
    ]
    for q in questions:
        lines.append(f"{q['number']}) [{q['type']}] {q['statement']}")
        for opt in q.get("options", []):
            lines.append(f"   {opt}")
        lines.append("")
    return "\n".join(lines).strip()


def _render_answer_key(questions: list[dict], topic: str) -> str:
    lines = [f"SOLUCIONARIO - {topic}", ""]
    for q in questions:
        lines.append(f"{q['number']}) Respuesta: {q['answer']}")
        if q.get("explanation"):
            lines.append(f"   Justificación: {q['explanation']}")
        lines.append("")
    return "\n".join(lines).strip()


def generate_exam(
    branch: str,
    topic: str,
    num_questions: int,
    difficulty: str,
    top_k: int,
    exam_type: Literal["test_simple", "test_multiple", "desarrollo", "mixto"] = "mixto",
) -> dict:
    if not branch_exists(branch):
        raise FileNotFoundError()

    contexts = retrieve_chunks(branch, topic, top_k)
    context_texts = [c["text"] for c in contexts]
    if not context_texts:
        raise RuntimeError("No hay suficiente contexto del temario para generar el examen.")

    prompt = (
        f"Genera un examen educativo sobre: {topic}. "
        f"Dificultad: {difficulty}. "
        f"Número de preguntas: {num_questions}. "
        f"{_build_type_instructions(exam_type)}"
    )
    context_block = _compact_context_texts(context_texts, max_chars=2200)
    full_prompt = (
        "Eres un profesor universitario experto. Crea preguntas válidas y no repetidas basadas solo en el contexto.\n"
        "Devuelve el resultado en este formato ESTRICTO por cada pregunta:\n"
        "### Pregunta N\n"
        "Tipo: test_simple | test_multiple | desarrollo\n"
        "Enunciado: ...\n"
        "Opciones:\n"
        "A) ...\nB) ...\nC) ...\nD) ...\n"
        "Respuesta: ...\n"
        "Explicacion: ...\n"
        "Para preguntas de desarrollo, deja la sección Opciones vacía.\n"
        "Para test_multiple, la Respuesta debe incluir varias letras separadas por coma (ej: A,C).\n"
        "No dejes ninguna pregunta sin respuesta.\n"
        "No repitas enunciados ni opciones casi idénticas.\n"
        "No incluyas texto fuera de ese formato.\n\n"
        f"Contexto:\n{context_block}\n\n"
        f"Instrucción: {prompt}"
    )
    raw_content = generate_text(full_prompt, max_tokens=900, temperature=0.2)
    questions = _parse_generated_questions(raw_content)
    if len(questions) < max(1, min(num_questions, 3)):
        repair_prompt = (
            "Reescribe el texto siguiente al formato exacto requerido. No inventes contenido nuevo.\n"
            "Formato exacto por bloque:\n"
            "### Pregunta N\n"
            "Tipo: test_simple | test_multiple | desarrollo\n"
            "Enunciado: ...\n"
            "Opciones:\n"
            "A) ...\nB) ...\nC) ...\nD) ...\n"
            "Respuesta: ...\n"
            "Explicacion: ...\n\n"
            f"Texto a reestructurar:\n{raw_content}\n"
        )
        repaired = generate_text(repair_prompt, max_tokens=900, temperature=0.1)
        questions = _parse_generated_questions(repaired)
        if not questions:
            questions = _parse_fallback_questions(raw_content)

    questions = _dedupe_questions(questions)

    if len(questions) < num_questions:
        missing = num_questions - len(questions)
        extra_questions = _complete_missing_questions(
            topic=topic,
            difficulty=difficulty,
            exam_type=exam_type,
            context_block=context_block,
            current_questions=questions,
            missing=missing,
        )
        questions = _dedupe_questions(questions + extra_questions)

    if len(questions) < max(1, min(num_questions, 3)):
        raise RuntimeError("No se pudo generar un examen estructurado. Intenta con otro temario o topic más concreto.")

    questions = questions[:num_questions]

    # --- Fase 4: validate distractors + confidence scoring ---
    questions = validate_exam_questions(questions, context_texts, use_llm=False)
    avg_confidence = round(
        sum(q.get("confidence", 1.0) for q in questions) / max(1, len(questions)), 2
    )
    distractor_warnings = [
        f"P{q['number']}: {issue}"
        for q in questions
        for issue in q.get("distractor_issues", [])
    ]

    exam_content = _render_exam_statement(questions, topic, difficulty, exam_type)
    answer_key_content = _render_answer_key(questions, topic)

    exams_dir = _exams_dir(branch)
    exams_dir.mkdir(parents=True, exist_ok=True)
    exam_id = uuid.uuid4().hex
    filename = f"exam_{exam_id}.json"
    payload = {
        "id": exam_id,
        "topic": topic,
        "difficulty": difficulty,
        "num_questions": num_questions,
        "exam_type": exam_type,
        "questions": questions,
        "exam_content": exam_content,
        "answer_key_content": answer_key_content,
        "raw_content": raw_content,
        "avg_confidence": avg_confidence,
        "distractor_warnings": distractor_warnings,
        "created_at": datetime.utcnow().isoformat(),
    }
    (exams_dir / filename).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "exam_id": exam_id,
        "filename": filename,
        "exam_content": exam_content,
        "answer_key_content": answer_key_content,
        "avg_confidence": avg_confidence,
        "distractor_warnings": distractor_warnings,
    }


def export_exam_pdf(branch: str, exam_id: str, kind: Literal["exam", "answer_key"] = "exam") -> Path:
    exam = _load_exam(branch, exam_id)
    exams_dir = _exams_dir(branch)
    suffix = "solucionario" if kind == "answer_key" else "enunciado"
    path = exams_dir / f"exam_{exam_id}_{suffix}.pdf"

    c = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    y = height - 72

    if kind == "answer_key":
        lines = exam.get("answer_key_content", "").splitlines()
    else:
        lines = exam.get("exam_content", exam.get("raw_content", "")).splitlines()

    for line in lines:
        if y < 72:
            c.showPage()
            y = height - 72
        c.drawString(72, y, line[:120])
        y -= 14

    c.save()
    return path


def export_exam_docx(branch: str, exam_id: str, kind: Literal["exam", "answer_key"] = "exam") -> Path:
    exam = _load_exam(branch, exam_id)
    exams_dir = _exams_dir(branch)
    suffix = "solucionario" if kind == "answer_key" else "enunciado"
    path = exams_dir / f"exam_{exam_id}_{suffix}.docx"

    doc = Document()
    if kind == "answer_key":
        doc.add_heading(f"Solucionario: {exam['topic']}", level=1)
        lines = exam.get("answer_key_content", "").splitlines()
    else:
        doc.add_heading(f"Examen: {exam['topic']}", level=1)
        doc.add_paragraph(f"Dificultad: {exam['difficulty']}")
        doc.add_paragraph(f"Tipo: {exam.get('exam_type', 'mixto')}")
        doc.add_paragraph("")
        lines = exam.get("exam_content", exam.get("raw_content", "")).splitlines()

    for line in lines:
        doc.add_paragraph(line)
    doc.save(str(path))
    return path


def solve_uploaded_exam(branch: str, file_path: Path, top_k: int = 8) -> dict:
    if not branch_exists(branch):
        raise FileNotFoundError()

    exam_text = extract_text_from_file(file_path)
    if not exam_text.strip():
        raise RuntimeError("No se pudo extraer texto del examen subido.")

    contexts = retrieve_chunks(branch, exam_text[:2000], top_k)
    context_block = _compact_context_texts([c["text"] for c in contexts], max_chars=3000)

    prompt = (
        "Resuelve el examen aportado.\n"
        "Devuelve SOLO el solucionario en formato:\n"
        "### Solución N\n"
        "Respuesta: ...\n"
        "Justificación: ...\n"
        "Si una pregunta es tipo test, indica opción(es) correctas.\n"
        "Si es de desarrollo, responde de forma estructurada y breve.\n"
        "Si falta información, dilo explícitamente en la justificación.\n\n"
        f"Contexto de apoyo del temario:\n{context_block}\n\n"
        f"Texto del examen a resolver:\n{exam_text}\n\n"
        "Solucionario:"
    )

    solutions = generate_text(prompt, max_tokens=1200, temperature=0.15)

    # Per-solution confidence scoring
    solution_blocks = re.split(r"###\s*Soluci[oó]n\s*\d+", solutions)
    context_texts_short = [c["text"] for c in contexts]
    solution_scores: list[float] = []
    for block in solution_blocks:
        block = block.strip()
        if block:
            score = score_solved_answer(block, context_texts_short)
            solution_scores.append(score)
    avg_solution_confidence = round(
        sum(solution_scores) / max(1, len(solution_scores)), 2
    ) if solution_scores else None

    return {"solutions": solutions, "avg_solution_confidence": avg_solution_confidence}


def start_exam_simulation(branch: str, exam_id: str, duration_minutes: int = 30) -> dict:
    exam = _load_exam(branch, exam_id)
    questions = exam.get("questions", [])
    if not questions:
        raise RuntimeError("El examen no contiene preguntas estructuradas para simulacro.")

    started_at_dt = datetime.utcnow()
    expires_at_dt = started_at_dt + timedelta(minutes=duration_minutes)
    simulation_id = uuid.uuid4().hex

    simulation_payload = {
        "simulation_id": simulation_id,
        "exam_id": exam_id,
        "topic": exam.get("topic", "Tema"),
        "duration_minutes": duration_minutes,
        "started_at": started_at_dt.isoformat(),
        "expires_at": expires_at_dt.isoformat(),
        "submitted_at": None,
        "status": "in_progress",
        "answers": [],
    }

    sim_dir = _simulations_dir(branch)
    sim_dir.mkdir(parents=True, exist_ok=True)
    (sim_dir / f"simulation_{simulation_id}.json").write_text(
        json.dumps(simulation_payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    return {
        "simulation_id": simulation_id,
        "exam_id": exam_id,
        "topic": exam.get("topic", "Tema"),
        "duration_minutes": duration_minutes,
        "started_at": simulation_payload["started_at"],
        "expires_at": simulation_payload["expires_at"],
        "questions": [
            {
                "number": int(q.get("number", i + 1)),
                "type": str(q.get("type", "desarrollo")),
                "statement": str(q.get("statement", "")),
                "options": list(q.get("options", [])),
            }
            for i, q in enumerate(questions)
        ],
    }


def _normalize_choice_letters(text: str) -> set[str]:
    letters = re.findall(r"\b([A-D])\b", text.upper())
    return set(letters)


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _development_is_correct(expected: str, student: str) -> bool:
    expected_tokens = set(re.findall(r"[a-záéíóúñ0-9]{4,}", _normalize_text(expected)))
    student_tokens = set(re.findall(r"[a-záéíóúñ0-9]{4,}", _normalize_text(student)))
    if not expected_tokens:
        return False
    overlap = len(expected_tokens.intersection(student_tokens))
    ratio = overlap / max(1, len(expected_tokens))
    return ratio >= 0.35


def _llm_grade_development(statement: str, expected: str, student: str) -> tuple[bool, str]:
    """
    Use the LLM to grade an open-ended development question.
    Returns (is_correct, feedback_text).
    Falls back to keyword heuristic if LLM call fails or returns unrecognisable output.
    """
    if not student.strip():
        return False, "Sin respuesta."

    prompt = (
        "Eres un corrector de exámenes. Evalúa si la respuesta del alumno es correcta.\n"
        f"Pregunta: {statement}\n"
        f"Respuesta esperada (modelo): {expected}\n"
        f"Respuesta del alumno: {student}\n\n"
        "Responde SOLO con el formato:\n"
        "VEREDICTO: CORRECTO o INCORRECTO\n"
        "FEEDBACK: <una frase corta de retroalimentación>\n"
    )
    try:
        raw = generate_text(prompt, max_tokens=80, temperature=0.0)
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        verdict = ""
        feedback = ""
        for line in lines:
            low = line.lower()
            if low.startswith("veredicto:"):
                verdict = line.split(":", 1)[1].strip().upper()
            elif low.startswith("feedback:"):
                feedback = line.split(":", 1)[1].strip()
        is_correct = "CORRECTO" in verdict and "INCORRECTO" not in verdict
        if not feedback:
            feedback = "Respuesta aceptable." if is_correct else "Revisa los conceptos clave."
        return is_correct, feedback
    except Exception:
        # Fall back to keyword heuristic
        heuristic = _development_is_correct(expected, student)
        feedback = (
            "Respuesta aceptable para desarrollo."
            if heuristic
            else "Respuesta incompleta para desarrollo; revisa conceptos clave."
        )
        return heuristic, feedback


def _question_topic_hint(statement: str) -> str:
    tokens = re.findall(r"[a-zA-ZáéíóúÁÉÍÓÚñÑ0-9]+", statement)
    return " ".join(tokens[:6]) if tokens else "tema general"


def submit_exam_simulation(branch: str, simulation_id: str, answers: list[dict]) -> dict:
    sim_path = _simulations_dir(branch) / f"simulation_{simulation_id}.json"
    if not sim_path.exists():
        raise FileNotFoundError()

    simulation = json.loads(sim_path.read_text(encoding="utf-8"))
    exam = _load_exam(branch, simulation["exam_id"])
    questions = exam.get("questions", [])
    if not questions:
        raise RuntimeError("El examen asociado no tiene preguntas estructuradas.")

    answer_map = {int(a.get("number", -1)): str(a.get("answer", "")).strip() for a in answers}
    results: list[dict] = []
    correct_count = 0
    answered_count = 0
    weak_topics: list[str] = []

    now = datetime.utcnow()
    expired = now > datetime.fromisoformat(simulation["expires_at"])

    for i, q in enumerate(questions, start=1):
        number = int(q.get("number", i))
        q_type = str(q.get("type", "desarrollo"))
        statement = str(q.get("statement", ""))
        expected = str(q.get("answer", "")).strip()
        student = answer_map.get(number, "")

        if student:
            answered_count += 1

        correct = False
        feedback = ""

        if not student:
            feedback = "Sin respuesta."
        elif q_type == "test_simple":
            expected_letters = _normalize_choice_letters(expected)
            student_letters = _normalize_choice_letters(student)
            if expected_letters and student_letters:
                correct = expected_letters == student_letters
            else:
                correct = _normalize_text(student) == _normalize_text(expected)
            feedback = "Correcta." if correct else "Revisa la opción correcta del temario."
        elif q_type == "test_multiple":
            expected_letters = _normalize_choice_letters(expected)
            student_letters = _normalize_choice_letters(student)
            correct = bool(expected_letters) and expected_letters == student_letters
            feedback = "Correcta." if correct else "Te faltan o sobran opciones correctas."
        else:
            correct, feedback = _llm_grade_development(statement, expected, student)

        if correct:
            correct_count += 1
        else:
            weak_topics.append(_question_topic_hint(statement))

        results.append(
            {
                "number": number,
                "type": q_type,
                "statement": statement,
                "student_answer": student,
                "expected_answer": expected,
                "correct": correct,
                "feedback": feedback,
                "question_confidence": float(q.get("confidence", 1.0)),
            }
        )

    total_questions = len(questions)
    score_percent = round((correct_count / max(1, total_questions)) * 100, 2)
    status = "timed_out" if expired else "submitted"

    simulation["answers"] = answers
    simulation["submitted_at"] = now.isoformat()
    simulation["status"] = status
    simulation["score_percent"] = score_percent
    sim_path.write_text(json.dumps(simulation, ensure_ascii=False, indent=2), encoding="utf-8")

    unique_weak_topics: list[str] = []
    seen: set[str] = set()
    for wt in weak_topics:
        k = _normalize_text(wt)
        if not k or k in seen:
            continue
        seen.add(k)
        unique_weak_topics.append(wt)

    avg_q_conf = round(
        sum(r.get("question_confidence", 1.0) for r in results) / max(1, len(results)), 2
    )

    return {
        "simulation_id": simulation_id,
        "exam_id": simulation["exam_id"],
        "topic": exam.get("topic", "Tema"),
        "total_questions": total_questions,
        "answered_questions": answered_count,
        "correct_answers": correct_count,
        "score_percent": score_percent,
        "status": status,
        "weak_topics": unique_weak_topics[:8],
        "results": results,
        "avg_question_confidence": avg_q_conf,
    }


def get_simulation_history(branch: str, limit: int = 30) -> dict:
    if not branch_exists(branch):
        raise FileNotFoundError()

    sim_dir = _simulations_dir(branch)
    if not sim_dir.exists():
        return {"items": []}

    items: list[dict] = []
    for path in sim_dir.glob("simulation_*.json"):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            items.append(
                {
                    "simulation_id": data.get("simulation_id", ""),
                    "exam_id": data.get("exam_id", ""),
                    "topic": data.get("topic", "Tema"),
                    "started_at": data.get("started_at", ""),
                    "submitted_at": data.get("submitted_at"),
                    "status": data.get("status", "in_progress"),
                    "score_percent": data.get("score_percent"),
                }
            )
        except Exception:
            continue

    items.sort(key=lambda x: (x.get("submitted_at") or x.get("started_at") or ""), reverse=True)
    return {"items": items[:limit]}


def _load_exam(branch: str, exam_id: str) -> dict:
    if not branch_exists(branch):
        raise FileNotFoundError()
    exams_dir = _exams_dir(branch)
    path = exams_dir / f"exam_{exam_id}.json"
    if not path.exists():
        raise FileNotFoundError()
    return json.loads(path.read_text(encoding="utf-8"))
